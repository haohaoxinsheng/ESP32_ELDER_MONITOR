from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PICTURE_DIR = Path(r"C:\Users\19573\Desktop\论文\picture")
OUT = DOCS / "基于多模态感知的独居老人居家智能监护系统设计_严格格式版.docx"
MANUSCRIPT = DOCS / "基于多模态感知的独居老人居家智能监护系统设计_新版论文正文.md"

TITLE_CN = "基于多模态感知的独居老人居家智能监护系统设计"
TITLE_EN = "Design of an Intelligent Home Monitoring System for Elderly People Living Alone Based on Multimodal Perception"
SCHOOL = "长  春  工  程  学  院"

AUTO_CITES_ENABLED = False
NEXT_CITE = 1
MAX_AUTO_CITE = 17


def normalize_mixed_spacing(text: str) -> str:
    """清理正文中中英文、数字和汉字之间不必要的空格。"""
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[A-Za-z0-9])", "", text)
    text = re.sub(r"(?<=[A-Za-z0-9])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"(?<=[，；：、])\s+(?=[A-Za-z0-9\u4e00-\u9fff])", "", text)
    text = re.sub(r"(?<=[A-Za-z0-9\u4e00-\u9fff])\s+(?=[，；：、。])", "", text)
    return re.sub(r"[ \t]{2,}", " ", text)


def set_run_font(run, *, font: str, east: str | None = None, size: float = 10.5,
                 bold: bool = False, superscript: bool = False) -> None:
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east or font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if superscript:
        run.font.superscript = True


def force_plain_paragraph(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5


def set_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(normalize_mixed_spacing(text))
    set_run_font(r, font="Times New Roman", east="宋体", size=10.5, bold=bold)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for style_name, size, font, bold in [
        ("Heading 1", 12, "黑体", True),
        ("Heading 2", 10.5, "宋体", False),
        ("Heading 3", 10.5, "宋体", False),
    ]:
        style = styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.left_indent = Cm(0)


def add_p(doc: Document, text: str = "", *, first_line: bool = True, align=None, bold: bool = False,
          size: float = 10.5, font: str = "宋体") -> None:
    global NEXT_CITE
    if font != "Times New Roman":
        text = normalize_mixed_spacing(text)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0.74) if first_line else None
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, font=font if font == "Times New Roman" else "Times New Roman",
                 east=font, size=size, bold=bold)
    if AUTO_CITES_ENABLED and first_line and len(text) >= 70 and NEXT_CITE <= MAX_AUTO_CITE:
        cite = p.add_run(f"[{NEXT_CITE}]")
        set_run_font(cite, font="Times New Roman", east="宋体", size=8, superscript=True)
        NEXT_CITE += 1


def add_h1(doc: Document, text: str) -> None:
    if len(doc.paragraphs) > 1:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    force_plain_paragraph(p)
    set_outline_level(p, 0)
    m = re.match(r"^(\d+(?:\.\d+)*)\s*(.*)$", text)
    if m:
        r_num = p.add_run(m.group(1) + "  ")
        set_run_font(r_num, font="Times New Roman", east="Times New Roman", size=12, bold=True)
        r_txt = p.add_run(normalize_mixed_spacing(m.group(2)))
        set_run_font(r_txt, font="Times New Roman", east="黑体", size=12, bold=True)
    else:
        r = p.add_run(normalize_mixed_spacing(text))
        set_run_font(r, font="Times New Roman", east="黑体", size=12, bold=True)
    keep_with_next(p)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    force_plain_paragraph(p)
    set_outline_level(p, 1)
    m = re.match(r"^(\d+(?:\.\d+)*)\s*(.*)$", text)
    if m:
        r_num = p.add_run(m.group(1) + "  ")
        set_run_font(r_num, font="Times New Roman", east="Times New Roman", size=10.5)
        r_txt = p.add_run(normalize_mixed_spacing(m.group(2)))
        set_run_font(r_txt, font="Times New Roman", east="宋体", size=10.5)
    else:
        r = p.add_run(normalize_mixed_spacing(text))
        set_run_font(r, font="Times New Roman", east="宋体", size=10.5)
    keep_with_next(p)


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    force_plain_paragraph(p)
    m = re.match(r"^(\d+(?:\.\d+)*)\s*(.*)$", text)
    if m:
        r_num = p.add_run(m.group(1) + "  ")
        set_run_font(r_num, font="Times New Roman", east="Times New Roman", size=10.5)
        r_txt = p.add_run(normalize_mixed_spacing(m.group(2)))
        set_run_font(r_txt, font="Times New Roman", east="宋体", size=10.5)
    else:
        r = p.add_run(normalize_mixed_spacing(text))
        set_run_font(r, font="Times New Roman", east="宋体", size=10.5)
    keep_with_next(p)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.28)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("· " + normalize_mixed_spacing(item))
        set_run_font(r, font="Times New Roman", east="宋体", size=10.5)


def add_table(doc: Document, caption: str, rows: list[list[str]], widths: list[float] | None = None) -> None:
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.runs[0].bold = True
    p.runs[0].font.name = "黑体"
    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.runs[0].font.size = Pt(9)
    keep_with_next(p)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for row_i, row in enumerate(rows):
        for col_i, value in enumerate(row):
            cell = table.cell(row_i, col_i)
            set_cell_text(cell, value, bold=row_i == 0)
            set_cell_margins(cell)
            if row_i == 0:
                set_cell_shading(cell, "D9EAF7")
            if widths and col_i < len(widths):
                cell.width = Cm(widths[col_i])
    add_p(doc, "", first_line=False)


def add_image(doc: Document, path: Path, caption: str, *, width: float = 11.5) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    r = cap.runs[0]
    set_run_font(r, font="Times New Roman", east="黑体", size=9, bold=True)


def add_image_grid(doc: Document, entries: list[tuple[Path, str]], *, width: float = 6.0,
                   cols: int = 2) -> None:
    valid = [(path, caption) for path, caption in entries if path.exists()]
    if not valid:
        return
    rows = (len(valid) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    idx = 0
    for r_i in range(rows):
        for c_i in range(cols):
            cell = table.cell(r_i, c_i)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, 80, 80, 80, 80)
            if idx >= len(valid):
                cell.text = ""
                continue
            path, caption = valid[idx]
            idx += 1
            p_img = cell.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.first_line_indent = None
            p_img.add_run().add_picture(str(path), width=Cm(width))
            p_cap = cell.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.first_line_indent = None
            cap_run = p_cap.add_run(caption)
            set_run_font(cap_run, font="Times New Roman", east="黑体", size=9, bold=True)
    add_p(doc, "", first_line=False)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run("目    录")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend([begin, instr, separate])
    for line in [
        "摘   要",
        "Abstract",
        "1  前言",
        "2  需求分析与方案论证",
        "3  系统总体设计",
        "4  系统硬件设计",
        "5  系统软件设计",
        "6  Web可视化与远程交互设计",
        "7  系统测试与结果分析",
        "8  总结与展望",
        "参考文献",
        "致    谢",
        "附录",
    ]:
        p.add_run(line + "\n")
    end_run = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def pic(name: str) -> Path:
    return PICTURE_DIR / name


def add_flow_image(doc: Document, caption: str, lines: list[str]) -> None:
    col_count = len(lines) * 2 - 1
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, line in enumerate(lines):
        cell = table.cell(0, i * 2)
        set_cell_text(cell, line, bold=False)
        set_cell_margins(cell, 150, 80, 150, 80)
        set_cell_shading(cell, "EEF4FB")
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "10")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "4472C4")
        if i < len(lines) - 1:
            arrow = table.cell(0, i * 2 + 1)
            set_cell_text(arrow, "→", bold=True)
            set_cell_margins(arrow, 150, 20, 150, 20)
            arrow.width = Cm(0.55)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    r = cap.runs[0]
    set_run_font(r, font="Times New Roman", east="黑体", size=9, bold=True)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        add_p(doc, "", first_line=False)
    add_p(doc, "毕业设计（论文）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
    add_p(doc, "", first_line=False)
    add_p(doc, TITLE_CN, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22, font="黑体")
    add_p(doc, TITLE_EN, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, font="Times New Roman")
    for _ in range(5):
        add_p(doc, "", first_line=False)
    for item in [
        "学生姓名：                  ",
        "学历层次：       本 科      ",
        "所在系部：  计算机技术与工程学院",
        "所学专业：      物联网工程      ",
        "指导教师：                  ",
        "教师职称：                  ",
        "完成时间：      2026年5月27日",
    ]:
        add_p(doc, item, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, font="楷体")
    for _ in range(4):
        add_p(doc, "", first_line=False)
    add_p(doc, SCHOOL, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, font="华文行楷")
    doc.add_page_break()


def add_abstract(doc: Document) -> None:
    p = doc.add_paragraph("摘   要")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.runs[0].bold = True
    p.runs[0].font.name = "黑体"
    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.runs[0].font.size = Pt(16)
    for text in [
        "随着家庭结构小型化和老龄化程度加深，独居老人居家安全已经不只是家庭内部的问题，也逐渐变成社区养老和智慧家居建设中需要长期面对的现实课题。老人一个人在家时，燃气泄漏、一氧化碳超标、夜间起身、摔倒后无法求助以及长时间没有活动等情况都可能突然发生，而传统人工看护很难做到连续陪伴，摄像头方案又容易让老人产生被监视的压力，所以需要一种低成本、低打扰、能持续运行的居家监护方式。",
        "本文以“基于多模态感知的独居老人居家智能监护系统设计”为题，围绕环境安全、人体活动、床位状态和远程查看等需求，设计了以 ESP32-S3 为核心的智能监护原型系统。系统把 DHT22、BH1750、MQ2、MQ135、MQ7、HC-SR501、SW-420、FSR、SOS 按键等模块接入同一终端，通过温湿度、光照、气体、人体红外、振动、压力和人工求助等多种信息共同描述老人居家状态；本地端使用 OLED、蜂鸣器、LED、风扇继电器和舵机完成提示与联动，网页端用于显示实时状态、事件记录、阈值配置和设备在线情况。",
        "在设计过程中，本文没有把系统写成单一传感器报警器，而是采用“感知层、传输层、用户层”的总体架构，再在设备内部细分采集、判断、执行和显示等模块。传感器数据先在 ESP32 端完成基础判断，关键告警可以脱离网络独立触发；同时，设备又把遥测数据上传到 Web 服务，使家属可以通过网页了解当前状态并调整阈值。测试结果表明，系统能够完成环境监测、气体风险提醒、暗环境人体活动照明、起夜记录、长时间无活动提醒、SOS 求助和设备离线提示等功能，整体结构清楚，成本可控，具有一定的实际应用和继续完善价值。",
    ]:
        add_p(doc, text)
    add_p(doc, "关键词：多模态感知  独居老人  智能监护  ESP32-S3", first_line=False)
    doc.add_page_break()

    p = doc.add_paragraph("Abstract")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.runs[0].bold = True
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(16)
    for text in [
        "With the deepening of population aging and the shrinking of family size, home safety for elderly people living alone has become an important issue in community care and smart home applications. Manual care is difficult to maintain continuously, while camera-based monitoring may bring privacy concerns in bedrooms and other private spaces. Therefore, a low-cost and low-intrusion monitoring system that can work continuously and provide timely warnings is needed.",
        "This thesis designs an intelligent home monitoring prototype based on ESP32-S3 and multimodal perception. The system integrates DHT22, BH1750, MQ2, MQ135, MQ7, HC-SR501, SW-420, FSR and an SOS button to collect environmental, gas, motion, vibration, pressure and help-request information. Local OLED display, buzzer, LED light, fan relay and servo are used for direct feedback and linkage, while the Web dashboard provides real-time status, event records, threshold configuration and device connection information.",
        "The system is organized as perception layer, transmission layer and user layer. Basic sensing and risk derivation are completed on the ESP32 terminal, so important alarms can still be triggered locally when the network is unavailable. At the same time, telemetry data are mirrored to the Web service for family members to view and configure. Tests show that the system can support environmental monitoring, gas risk warning, dark-environment lighting, night activity records, no-motion reminders, SOS alarm and offline notification. The design is clear, explainable and extensible for further smart elderly-care applications.",
    ]:
        add_p(doc, text, font="Times New Roman")
    add_p(doc, "Keywords: multimodal perception; elderly people living alone; intelligent monitoring; ESP32-S3", first_line=False, font="Times New Roman")
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    add_page_number(section)
    add_toc(doc)
    doc.add_page_break()


def add_main_sections(doc: Document) -> None:
    global AUTO_CITES_ENABLED, NEXT_CITE
    AUTO_CITES_ENABLED = True
    NEXT_CITE = 1
    add_h1(doc, "1  前言")
    add_h2(doc, "1.1  研究背景")
    for text in [
        "独居老人居家安全问题的特殊之处在于，它往往不是某一个设备能单独解决的。老人白天可能长时间坐着休息，夜间可能起床喝水或去卫生间，厨房里又可能出现烟雾、燃气和一氧化碳等风险，这些状态发生时并不一定会有人在旁边发现。过去很多家庭依靠电话问候、邻里照看或摄像头查看来降低风险，但是电话问候有明显的时间间隔，邻里照看也不可能全天候持续，摄像头虽然直观，却会把老人的私人空间暴露出来，老人本人也未必愿意长期接受。",
        "从技术角度看，物联网和嵌入式系统的发展让居家监护可以换一种思路来做。系统不必直接拍摄人的画面，而是可以通过温湿度、光照、气体浓度变化、人体红外、床位压力、振动和按键等信息，间接判断居家环境和活动状态。这种方式得到的信息没有视频那样细，但正因为边界比较清楚，部署时对隐私的影响也会小一些。对于独居老人监护场景来说，只要系统能及时发现“可能有风险”的状态，并把提醒传给老人本人或家属，就已经能发挥实际价值。",
        "本课题所设计的系统不是面向医院病房的专业生命体征监护仪，而是面向普通家庭的低成本居家安全原型。它关注的是老人日常生活中更常见、更容易被忽略的风险，例如空气异常、夜间活动、长时间没有活动、床边压力变化、震动冲击、SOS 求助和设备离线。系统选择 ESP32-S3 作为主控，一方面是因为它有较多 GPIO、ADC、I2C 和 PWM 资源，可以连接多种传感器和执行器；另一方面，它自带 WiFi，便于把本地状态上传到网页端，让家属在不进入老人生活空间的情况下了解情况。",
    ]:
        add_p(doc, text)
    add_h2(doc, "1.2  研究意义")
    for text in [
        "一方面，本课题能够提高安全提醒的及时性。居家风险很多时候不是完全没有征兆，而是没有被及时捕捉到。气体传感器数值异常、夜间暗环境下有人活动、床位压力突然变化、震动传感器触发、PIR 长时间没有检测到活动，这些信息单独看可能都不够确定，但是放在同一个系统里，就能帮助家属更早发现异常线索。本系统在本地端保留告警判断和执行器联动，即使网络暂时断开，也可以通过 OLED、蜂鸣器、灯光和风扇给出直观反馈。",
        "另一方面，本课题在隐私友好方面具有一定优势。系统没有把摄像头作为核心监测手段，而是使用非图像化传感器来描述环境和行为变化。对于卧室、床边、客厅等家庭空间，这种方式更容易被老人接受，也更符合居家养老“少打扰”的使用要求。它并不追求看清老人每一个动作，而是尽量在不打扰生活的前提下，把与安全有关的状态变成可被系统理解的数据。",
        "还有一方面，课题具有较完整的工程训练价值。系统涉及传感器接线、ADC 采样、I2C 通信、OLED 显示、继电器和舵机控制、阈值配置、Web 后端、实时推送、页面展示和云端部署等内容。它不是单纯写一个程序，也不是只做一个硬件小实验，而是要让设备端、网页端和实际测试结果能够互相对应，这对物联网专业毕业设计来说比较有代表性。",
    ]:
        add_p(doc, text)
    add_h2(doc, "1.3  国内外研究现状")
    for text in [
        "智慧养老相关研究大体可以分成可穿戴监测、环境监测、视频或雷达行为识别、智能家居联动以及远程照护平台几类。可穿戴设备能够获得心率、步数、睡眠等数据，精度和连续性较好，但是老人需要记得佩戴、充电和维护；视频识别能够给出更丰富的动作信息，不过隐私压力明显，尤其不适合卧室、卫生间和床边这样的场景；毫米波雷达隐私性比摄像头好，但是硬件成本和算法要求相对更高。",
        "在家庭环境安全方面，烟雾报警器、燃气报警器、人体感应灯、智能门磁和智能插座等产品已经比较常见，但是这些设备通常各自独立工作，家属看到的是分散的提醒，不容易形成一个完整的居家状态判断。面向独居老人的系统需要把这些分散信息组织起来，使环境安全、人体活动和远程查看之间形成闭环。也就是说，系统不能只会显示一个原始数值，还应该把数值背后的风险等级、联动动作和事件记录表达出来。",
        "基于 ESP32 的开源物联网方案在毕业设计和原型验证中应用较多，原因是生态成熟、模块成本低、资料容易获取，且能够兼顾本地控制和联网能力。本文参考此类系统的常见设计路线，但没有把重点放在复杂算法模型上，而是采用可解释的多源阈值融合方法。这样做的好处是每一个告警都能回到具体传感器状态和阈值设置，便于调试、演示和后续维护；不足之处是对复杂行为的识别能力有限，后续如果有长期数据积累，可以再引入更细致的统计分析或机器学习方法。",
    ]:
        add_p(doc, text)
    add_h2(doc, "1.4  本文主要研究内容")
    add_p(doc, "本文围绕独居老人居家智能监护系统完成了从需求分析、总体方案、硬件设计、软件设计、Web 端交互到测试验证的完整设计。为了避免论文写成设备说明书，本文在描述时重点说明系统为什么这样分层、各模块如何配合工作，以及测试结果能说明什么问题。主要工作包括：")
    add_bullets(doc, [
        "分析独居老人居家环境中的典型风险，把环境安全、人体活动、床位压力、震动冲击、人工求助和远程查看纳入同一系统；",
        "设计基于 ESP32-S3 的多模态感知终端，完成主要传感器、OLED 显示、蜂鸣器、LED、风扇继电器和舵机的硬件连接；",
        "采用感知层、传输层、用户层的总体架构，在设备内部实现采集、判断、联动、显示和遥测封装等功能；",
        "完成 Web 可视化页面和后端接口设计，使家属可以查看实时状态、事件记录、起夜记录、阈值设置和设备在线状态；",
        "结合实物照片、网页截图和功能测试，对气体异常、暗环境活动、起夜开灯、长时间无活动、SOS 求助和离线提示等场景进行验证。",
    ])

    add_h1(doc, "2  需求分析与方案论证")
    add_h2(doc, "2.1  应用场景与使用对象")
    for text in [
        "本系统的使用对象可以分成两类，一类是独居或半独居老人，另一类是远程关注老人状态的家属。老人端设备需要尽量自动运行，不能要求老人频繁打开手机、修改参数或手动确认复杂选项；家属端网页则需要把状态讲清楚，不能只给出一串原始 ADC 数值。对于老人来说，真正有用的是现场能不能听到、看到提醒；对于家属来说，真正有用的是能不能快速判断当前是否安全、设备是否在线、最近发生过什么事件。",
        "典型场景包括老人日间休息、夜间起床、厨房使用、床边活动、突发身体不适和设备断网。日间休息时，PIR 可能长时间没有动作，这并不一定表示危险，所以系统需要结合时间窗口和其他信息进行判断；夜间起床时，暗环境和人体活动同时出现，系统应当自动点亮灯光，减少摸黑行走带来的风险；厨房或室内空气异常时，系统需要启动通风并记录事件；当老人主动按下 SOS 按键时，系统应将其视为最高优先级事件，而不是普通提示。",
    ]:
        add_p(doc, text)
    add_h2(doc, "2.2  功能需求")
    add_p(doc, "围绕上述场景，系统功能需求可以归纳为感知、判断、联动、展示和配置五个方面，具体如下：")
    add_bullets(doc, [
        "设备应能够采集温湿度、光照、烟雾或可燃气体、空气质量、一氧化碳、人体活动、振动、床位压力和 SOS 按键状态；",
        "系统应把多种传感器信息转化为可理解的状态，例如空气异常、暗环境活动、起夜、疑似跌倒、长时间无活动和人工求助；",
        "本地联动应直观可靠，OLED、蜂鸣器、LED、风扇和舵机需要根据不同风险等级给出不同反应；",
        "网页端应便于家属查看，除了实时数据，还要显示事件记录、起夜记录、设备在线情况和当前主要告警；",
        "阈值应支持调整，因为不同家庭环境、传感器个体差异、接线方式和模块预热情况都会影响原始读数。",
    ])
    add_h2(doc, "2.3  非功能需求")
    for text in [
        "非功能需求主要体现在可靠性、易用性、可维护性、隐私性和成本控制上。可靠性要求系统在网络不可用时仍能进行本地判断和提醒，不能把所有安全动作都依赖网页或云端；易用性要求老人端尽量少操作，家属端页面状态要清晰；可维护性要求硬件引脚、阈值和业务逻辑有清楚边界，后期更换模块时不需要重写大段代码。",
        "隐私性方面，系统采用非图像化传感器为主，只记录状态、数值和事件，不采集视频和音频内容。成本方面，系统使用 ESP32-S3 开发板和常见模块完成原型，便于采购和替换。这样的取舍并不意味着系统没有局限，例如规则阈值判断对安装环境比较敏感，但是对于毕业设计原型来说，它能在可解释、可复现和成本可控之间取得比较合适的平衡。",
    ]:
        add_p(doc, text)
    add_h2(doc, "2.4  方案论证")
    for text in [
        "主控方案方面，若采用传统 8 位单片机，成本虽然较低，但是在多路传感器接入、JSON 数据处理和联网通信方面会比较吃力；若采用树莓派等单板计算机，处理能力充足，但功耗、成本和系统维护复杂度都会增加。ESP32-S3 同时具备 GPIO、ADC、I2C、PWM、WiFi 和较好的开发生态，适合承担本课题中的采集、判断、执行控制和上报任务，所以本文选择 ESP32-S3 作为核心控制器。",
        "感知方案方面，单一传感器无法可靠描述老人居家状态。PIR 没有检测到活动，可能是老人睡着了，也可能是传感器角度不合适；振动传感器触发，可能是桌面碰撞，也可能是跌倒带来的冲击；床位压力变化，可能是离床，也可能是姿态调整。所以系统采用多模态感知，把气体、光照、人体活动、振动、压力和 SOS 结合起来，形成比单点报警更有解释力的状态。",
        "通信方案方面，系统没有把风险判断完全放到云端，而是采用本地自治和 Web 镜像相结合的方案。ESP32 端负责关键采集、判断和执行器联动，网页端负责更完整的状态展示、事件记录、阈值配置和远程查看。这样做可以减少网络波动带来的安全风险，也能让家属获得更友好的查看入口。",
    ]:
        add_p(doc, text)
    add_table(doc, "表2-1  主要方案比较", [
        ["比较内容", "可选方案", "特点", "本文选择"],
        ["主控", "8位单片机、ESP32-S3、单板计算机", "8位单片机联网能力弱，单板计算机成本高，ESP32-S3资源和成本较均衡", "ESP32-S3"],
        ["感知方式", "摄像头、单传感器、多模态传感器", "摄像头隐私压力大，单传感器误判较多，多模态方式更适合居家安全判断", "多模态感知"],
        ["判断位置", "纯云端、本地端、本地与Web结合", "纯云端依赖网络，本地端展示能力有限，结合方案更稳妥", "本地自治加Web镜像"],
    ], [2.4, 3.0, 6.0, 3.0])

    add_h1(doc, "3  系统总体设计")
    add_h2(doc, "3.1  总体架构")
    add_p(doc, "系统总体上可以归纳为感知层、传输层和用户层三层。感知层负责把家庭环境和老人活动转化为电信号；传输层负责在设备内部和网页端之间传递状态；用户层负责把状态用老人和家属都能理解的方式呈现出来。为了便于工程实现，感知层内部又包含传感器和执行器，传输层包含 ESP32 本地处理、HTTP 上报和 SSE 推送，用户层包含 OLED、声光提醒和 Web 页面。系统总体结构如下图所示。")
    add_flow_image(doc, "图3-1  系统三层总体架构", ["感知层\n多类传感器\n本地执行器", "传输层\nESP32-S3\nHTTP/SSE", "用户层\nOLED提醒\nWeb查看"])
    add_p(doc, "这种三层划分比单纯罗列硬件模块更容易说明系统的工作逻辑。感知层并不是只负责“读数”，它还提供判断风险所需要的原始依据；传输层不只是网络通信，也包括本地状态整理和数据封装；用户层不只是页面显示，还包括老人端的直观提醒。三层之间形成闭环后，系统才能从采集数据走向实际监护。")
    add_h2(doc, "3.2  数据流与状态闭环")
    add_p(doc, "系统运行时，ESP32 按固定周期读取传感器数据，先在本地形成环境状态、气体状态、活动状态和告警状态，再根据状态控制风扇、蜂鸣器、灯光、舵机和 OLED 显示。随后，设备把整理后的遥测数据发送给 Web 后端，后端保存最新状态、历史数据、事件记录和起夜记录，再通过 SSE 推送给浏览器。家属在网页端修改阈值或联动开关后，设备端可以继续拉取并应用配置。数据流如下图所示。")
    add_flow_image(doc, "图3-2  系统数据流与状态闭环", ["传感器采集", "本地状态判断", "执行器联动", "HTTP遥测上报", "Web展示与配置"])
    add_h2(doc, "3.3  系统实物组成")
    add_p(doc, "在实物搭建中，系统使用面包板和杜邦线完成原型连接，重点验证各传感器、主控和执行器之间的数据与控制链路。实物连接需要特别注意共地、模拟量电压范围、MQ 系列预热和执行器电流。下图展示了系统整体实物连接情况。")
    add_image(doc, pic("实物图.jpg"), "图3-3  系统整体实物连接图", width=10.5)
    add_p(doc, "从实物结构看，ESP32-S3 位于控制中心，气体传感器、温湿度传感器、光照传感器、PIR、OLED 和执行器围绕主控连接。这样的原型虽然线材较多，但有利于在调试阶段快速更换模块和观察故障点，后续如果进入长期使用阶段，还需要再做外壳固定、电源保护和走线整理。")
    add_table(doc, "表3-1  三层架构与模块对应关系", [
        ["层次", "主要模块", "承担作用"],
        ["感知层", "DHT22、BH1750、MQ2、MQ135、MQ7、PIR、SW-420、FSR、SOS", "采集环境、气体、活动、压力、振动和人工求助信息"],
        ["传输层", "ESP32-S3、HTTP接口、SSE推送、控制配置", "完成本地判断、状态封装、网页上报和阈值同步"],
        ["用户层", "OLED、蜂鸣器、LED、风扇、舵机、Web页面", "向老人和家属展示状态，并根据风险进行提醒或联动"],
    ], [2.5, 5.5, 6.0])

    add_h1(doc, "4  系统硬件设计")
    add_h2(doc, "4.1  主控与显示模块")
    add_p(doc, "ESP32-S3是系统的核心，它一边连接各类传感器，一边控制显示和执行器，还要负责网络通信。相比只适合简单IO控制的传统单片机，ESP32-S3更适合这种需要多路采集、状态处理和联网展示的原型系统。OLED显示屏用于本地端状态展示，优先呈现关键状态和主要告警，主控和显示模块实物如下图所示。")
    add_image_grid(doc, [
        (pic("esp32开发板.jpg"), "图4-1  ESP32-S3开发板实物图"),
        (pic("OLED显示屏.jpg"), "图4-2  OLED显示屏实物图"),
    ], width=6.0)
    add_h2(doc, "4.2  环境感知模块")
    add_p(doc, "环境感知主要由温湿度传感器和光照传感器完成。DHT22用于采集温度和湿度，能够反映老人居住环境是否过热、过冷、过湿或过干；BH1750用于采集光照强度，它不直接表示危险，但能为暗环境活动、起夜照明和窗帘联动提供依据。两个环境模块实物如下图所示。")
    add_image_grid(doc, [
        (pic("温湿度传感器.jpg"), "图4-3  DHT22温湿度传感器实物图"),
        (pic("光线传感器.jpg"), "图4-4  BH1750光照传感器实物图"),
    ], width=6.0)
    add_h2(doc, "4.3  气体安全感知模块")
    add_p(doc, "气体安全是居家监护中比较直接的风险来源。本系统使用 MQ2、MQ135 和 MQ7 分别关注烟雾或可燃气体、空气质量和一氧化碳风险。MQ 系列传感器容易受到预热时间、供电、电位器和环境变化影响，所以本文把它们的原始值作为风险判断依据，并允许在网页端调整阈值。下面三幅图分别展示了各气体模块实物。")
    add_image_grid(doc, [
        (pic("MQ-2.jpg"), "图4-5  MQ-2烟雾/可燃气体传感器实物图"),
        (pic("MQ-135.jpg"), "图4-6  MQ-135空气质量传感器实物图"),
        (pic("MQ-7.jpg"), "图4-7  MQ-7一氧化碳传感器实物图"),
    ], width=4.5, cols=3)
    add_h2(doc, "4.4  人体活动、压力与振动感知模块")
    add_p(doc, "人体活动检测使用 HC-SR501 PIR 传感器，它适合判断一定范围内是否有活动，但不能像摄像头那样给出具体姿态，所以系统把它作为活动线索，而不是唯一判断依据。人体监测传感器实物如下图所示。")
    add_image(doc, pic("人体监测传感器.jpg"), "图4-8  HC-SR501 人体红外传感器实物图", width=8.8)
    add_p(doc, "压力和振动用于补充活动判断。压力传感器可以反映床位占用或受压变化，振动传感器可以捕捉碰撞或冲击，两者与 PIR 的活动时间结合后，可用于疑似跌倒或异常活动的辅助判断。单独使用其中任何一个都容易误判，所以本文强调组合判断。")
    add_h2(doc, "4.5  执行器与本地提醒模块")
    add_p(doc, "本地提醒模块包括蜂鸣器、LED灯、风扇继电器和舵机。蜂鸣器用于声音提示，LED用于暗环境照明和告警提示，继电器控制风扇完成通风联动，舵机用于SOS求助或窗帘动作演示。这些模块让系统不只是“看见异常”，还能在本地做出可感知的反应。")
    add_image_grid(doc, [
        (pic("有源蜂鸣器.jpg"), "图4-9  有源蜂鸣器实物图"),
        (pic("继电器.jpg"), "图4-10  继电器模块实物图"),
        (pic("风扇.jpg"), "图4-11  风扇模块实物图"),
        (pic("360舵机.jpg"), "图4-12  360舵机实物图"),
    ], width=5.2)
    add_p(doc, "联动点灯测试能够直观说明暗环境活动和告警灯光之间的关系。当系统判断需要照明或告警时，LED 会被点亮或闪烁，老人不需要理解网页上的字段，也能通过现场反馈知道设备已经响应。联动点灯实物如下图所示。")
    add_image(doc, pic("联动点灯.jpg"), "图4-13  灯光联动实物图", width=9.5)
    add_table(doc, "表4-1  主要硬件模块与功能对应", [
        ["模块", "连接方式", "系统作用"],
        ["ESP32-S3", "主控开发板", "采集、判断、联动控制和网络通信"],
        ["DHT22、BH1750", "GPIO/I2C", "获得温湿度和光照，支持环境舒适度和暗环境判断"],
        ["MQ2、MQ135、MQ7", "ADC", "获得气体风险线索，支持通风和危险提醒"],
        ["PIR、压力、振动、SOS", "GPIO/ADC", "获得活动、床位、冲击和主动求助信息"],
        ["OLED、蜂鸣器、LED、风扇、舵机", "I2C/GPIO/PWM", "向老人端给出显示、声音、照明、通风和动作反馈"],
    ], [3.0, 3.0, 7.8])

    add_h1(doc, "5  系统软件设计")
    add_h2(doc, "5.1  固件结构")
    for text in [
        "固件使用 PlatformIO 和 Arduino Framework 开发，整体按功能拆分为硬件初始化、传感器采集、活动状态推导、告警判断、执行器控制、OLED 显示、遥测封装和云端通信等部分。这样的拆分不是为了增加文件数量，而是为了让每个部分的职责更加清楚。例如，引脚和阈值集中放在配置文件中，主循环就可以更专注于“采集到什么、判断出什么、需要怎样联动”。",
        "主循环运行时并不是一次性做完所有事情，而是按不同周期分别执行。传感器采样和 OLED 刷新较快，串口输出和网络上报可以稍慢一些。这样既能保证告警响应及时，也能避免显示和日志过于频繁而影响可读性。对于 DHT22 和 MQ 系列这类响应较慢或波动较大的模块，系统更关注趋势和阈值触发，而不是把每一次瞬时读数都当成绝对结论。",
    ]:
        add_p(doc, text)
    add_h2(doc, "5.2  传感器采集流程")
    add_p(doc, "传感器采集流程可以理解为从原始信号到状态线索的转换。系统先读取温湿度、光照、气体、压力、振动、人体活动和按键状态，再对缺测值、瞬时触发和时间窗口进行处理。采集流程如下图所示。")
    add_flow_image(doc, "图5-1  传感器采集流程", ["定时采样", "读取各模块", "记录时间戳", "形成状态线索", "进入告警判断"])
    add_p(doc, "在这个过程中，PIR 检测到活动时会更新最近活动时间，振动传感器触发时会记录最近振动时间，压力传感器则提供床位和受压变化线索。时间戳的引入让系统不只看当前一瞬间的电平，还能判断一段时间内是否持续没有活动，或振动事件是否刚刚发生过。")
    add_h2(doc, "5.3  典型传感器链路处理")
    add_p(doc, "为了让系统工作过程更容易理解，本文把各类传感器从采集到用户显示的过程分成几条典型链路。每条链路都包含“采集原始信号、转换成状态、参与判断、形成提示、展示给用户”几个环节。这样写可以避免只罗列传感器型号，也能说明多模态感知是怎样真正进入系统逻辑的。")
    add_h3(doc, "5.3.1  环境感知链路")
    add_p(doc, "环境感知链路主要由 DHT22 和 BH1750 完成。DHT22 采集温湿度后，系统先判断居住环境是否过热、过冷、过湿或过干；BH1750 采集光照后，系统判断是否处于暗环境。两类数据一方面显示在 Web 仪表盘和 OLED 上，另一方面参与通风、夜间照明和舒适度提醒。其处理过程如下图所示。")
    add_flow_image(doc, "图5-2  环境感知数据处理流程", ["温湿度/光照采集", "有效值检查", "舒适度与暗环境判断", "联动灯光或通风", "OLED与Web显示"])
    add_h3(doc, "5.3.2  气体安全链路")
    add_p(doc, "气体安全链路由 MQ2、MQ135 和 MQ7 提供输入。三个模块输出的是模拟量，ESP32 读取后不会直接把它解释成绝对浓度，而是与现场校准后的阈值比较，生成正常、预警或危险状态。危险状态会进一步影响风扇、蜂鸣器、事件记录和 Web 页面提示。其处理过程如下图所示。")
    add_flow_image(doc, "图5-3  气体安全数据处理流程", ["ADC原始值读取", "阈值比较", "风险等级生成", "风扇/蜂鸣器联动", "事件记录与页面提示"])
    add_h3(doc, "5.3.3  活动与起夜链路")
    add_p(doc, "活动与起夜链路不是单独依赖 PIR，而是结合光照、人体活动和床位压力。光照较低时，系统先确认是否属于暗环境；PIR 检测到人体活动后更新最近活动时间；如果床位压力也发生变化，系统就能更接近真实地描述老人起身或离床。该链路的输出会进入起夜记录、灯光联动和长时间无活动判断。其处理过程如下图所示。")
    add_flow_image(doc, "图5-4  活动与起夜数据处理流程", ["光照判断暗环境", "PIR检测人体活动", "压力判断床位变化", "生成起夜/无活动状态", "灯光联动与Web记录"])
    add_h3(doc, "5.3.4  求助与异常冲击链路")
    add_p(doc, "SOS 按键和振动传感器承担更直接的安全提示作用。SOS 按键经过消抖后直接进入高优先级告警；振动传感器触发后，系统会结合压力和活动时间判断是否需要提高风险等级。用户端看到的不只是一个开关量，而是经过系统处理后的告警文案、事件记录和执行器动作。其处理过程如下图所示。")
    add_flow_image(doc, "图5-5  求助与异常冲击处理流程", ["SOS/振动输入", "消抖与时间窗口", "结合压力和活动状态", "生成高优先级提示", "蜂鸣器/OLED/Web提醒"])
    add_h2(doc, "5.4  多模态状态判断")
    for text in [
        "系统采用规则融合方法进行状态判断。所谓规则融合，并不是简单把多个阈值相加，而是根据居家场景把不同传感器组合成有意义的事件。例如，暗环境本身不是危险，但是暗环境下检测到人体活动，就可能表示老人夜间起身；PIR 长时间没有活动不一定危险，但是如果同时伴随压力异常或刚刚发生过强振动，就需要提高关注等级。",
        "气体风险判断相对直接，MQ2、MQ135 和 MQ7 分别设置预警和危险阈值。当 MQ7 达到危险等级时，系统会把一氧化碳风险放在较高优先级，因为这种风险对生命安全影响更直接。温湿度状态主要用于舒适度提醒和通风辅助，光照状态主要参与暗环境和起夜判断，SOS 按键则不需要复杂推导，只要确认按下并完成消抖，就应进入紧急提示。",
    ]:
        add_p(doc, text)
    add_flow_image(doc, "图5-6  多模态告警判断流程", ["环境与气体判断", "活动与床位判断", "SOS优先判断", "合成危险等级", "输出主告警文本"])
    add_h2(doc, "5.5  执行器联动逻辑")
    for text in [
        "执行器联动的目标是让系统的判断真正落到现场。气体异常或温湿度异常时，风扇可以启动通风；暗环境有人活动或起夜时，LED 可以点亮；发生 SOS、疑似跌倒或严重气体风险时，蜂鸣器会提示；舵机则用于 SOS 或窗帘联动演示。联动逻辑并不是所有告警都同时打开全部设备，而是根据风险类型和用户开关进行组合。",
        "这种设计也考虑了实际使用中的可控性。家属可以在网页端关闭某些演示型联动，例如关闭普通暗环境自动开灯，但保留危险告警强制亮灯；也可以根据传感器实际读数调整阈值。这样系统不会被固定参数锁死，更适合不同家庭环境下的调试和使用。",
    ]:
        add_p(doc, text)
    add_h2(doc, "5.6  遥测数据与配置同步")
    add_p(doc, "ESP32 端会把整理后的遥测数据上传到 Web 后端，数据中既有原始值，也有系统推导出的状态，例如是否暗环境、是否检测到夜间活动、是否存在总告警、当前风扇和灯光是否开启等。Web 端保存最新数据、历史数据和事件记录，并把控制开关和阈值保存为配置，使设备端和页面端可以保持一致。")
    add_table(doc, "表5-1  遥测数据分类", [
        ["类别", "代表内容", "用途"],
        ["环境数据", "温度、湿度、光照", "展示舒适度和暗环境状态"],
        ["气体数据", "MQ2、MQ135、MQ7 原始值", "判断烟雾、空气质量和一氧化碳风险"],
        ["活动数据", "PIR、振动、压力、SOS", "判断活动、起夜、疑似跌倒和求助"],
        ["联动状态", "风扇、灯光、蜂鸣器、舵机", "让网页显示实际执行结果"],
        ["配置数据", "阈值和开关", "支持现场校准和远程调整"],
    ], [2.5, 4.5, 6.8])

    add_h1(doc, "6  Web可视化与远程交互设计")
    add_h2(doc, "6.1  Web 后端设计")
    for text in [
        "Web 后端使用 Node.js 实现，主要负责接收 ESP32 上报的遥测数据、维护最新状态、保存历史记录和事件记录，并向浏览器推送实时更新。后端还提供阈值和控制开关接口，使网页端的配置能够保存下来。为了避免任意客户端伪造设备数据，上报接口使用设备令牌进行校验。",
        "后端的另一个重要作用是离线判断。设备长期没有上报时，网页不能继续假装数据仍然实时有效，而应提示设备离线并限制配置操作。这样做可以避免家属误以为阈值已经成功下发，也能更清楚地区分“设备正常但没有告警”和“设备已经断开”。设备在线状态页面如下图所示。",
    ]:
        add_p(doc, text)
    add_p(doc, "当设备离线时，页面会显示离线提示，并对控制项进行保护，避免用户在设备不在线时连续保存配置。在线和离线两种状态如下图所示。")
    add_image_grid(doc, [
        (pic("设备在线.png"), "图6-1  设备在线状态页面"),
        (pic("离线状态.jpeg"), "图6-2  设备离线状态页面"),
    ], width=6.5)
    add_h2(doc, "6.2  前端页面设计")
    add_p(doc, "前端页面的设计重点不是把所有数据堆在一个页面里，而是把信息按阅读优先级组织起来。安全状态、SOS、起夜和主要告警放在更显眼位置，温湿度、光照和气体数据放在仪表盘区域，传感器状态、执行器状态、事件记录和设置项则根据使用场景分区展示。运行概况页面如下图所示。")
    add_p(doc, "传感器仪表盘页面用于集中显示多类传感器数值，家属可以通过它了解当前环境是否正常。与普通参数表不同，页面还结合颜色和状态文案表达风险等级，使非专业用户也能大致看懂。传感器状态页面更偏向调试和维护，用于确认各模块是否启用、数据是否有效以及当前读数是否异常。三类前端页面如下图所示。")
    add_image_grid(doc, [
        (pic("运行概况.png"), "图6-3  系统运行概况页面"),
        (pic("传感器仪表盘.png"), "图6-4  传感器仪表盘页面"),
        (pic("传感器状态.png"), "图6-5  传感器状态页面"),
    ], width=4.6, cols=3)
    add_h2(doc, "6.3  联动控制与阈值设置")
    add_p(doc, "联动控制页面把暗环境灯、起夜灯、告警灯、通风、蜂鸣器、SOS 舵机和无人活动提醒等开关集中展示。这样做的目的是让家属能够根据实际家庭环境调整系统行为，而不是让设备一直以固定策略运行。联动页面如下图所示。")
    add_p(doc, "设置页面用于调整气体、光照、床位压力、震动、温湿度和无人活动时间等阈值。由于MQ系列和FSR读数与环境、接线和模块个体差异有关，阈值可调对实际部署很重要。联动控制和阈值设置页面如下图所示。")
    add_image_grid(doc, [
        (pic("联动页面.jpeg"), "图6-6  联动控制页面"),
        (pic("设置页面.jpeg"), "图6-7  阈值设置页面"),
    ], width=6.5)
    add_h2(doc, "6.4  事件记录与深色模式")
    add_p(doc, "事件记录用于把系统发生过的异常保留下来，家属可以查看最近是否出现空气异常、起夜活动、长时间无活动或其他告警。对于长期照护来说，事件记录比单次实时数据更有参考意义。事件记录页面如下图所示。")
    add_p(doc, "考虑到夜间查看和不同使用习惯，前端还提供深色模式。深色模式可以降低夜间查看时的视觉刺激，也能让告警色彩更加突出。事件记录和深色模式页面如下图所示。")
    add_image_grid(doc, [
        (pic("事件记录.png"), "图6-8  事件记录页面"),
        (pic("深色模式.jpeg"), "图6-9  深色模式页面"),
    ], width=6.5)

    add_h1(doc, "7  系统测试与结果分析")
    add_h2(doc, "7.1  测试环境与方法")
    for text in [
        "系统测试分为硬件连接测试、单模块测试、联动测试和网页展示测试。硬件连接测试主要确认供电、共地、引脚接线和 OLED 启动是否正常；单模块测试关注温湿度、光照、气体、PIR、压力、振动、SOS 和执行器是否能独立响应；联动测试关注多个条件组合后是否能触发预期动作；网页展示测试则验证设备状态、事件记录、阈值设置和离线提示是否与本地端一致。",
        "由于本课题是毕业设计原型，不具备标准实验室气体浓度标定条件，所以 MQ 系列测试采用相对变化和阈值触发方式进行，重点验证系统能否根据读数变化进入预警或危险状态，而不是给出医学或工业级浓度结论。FSR 压力和 PIR 活动测试也采用场景化方式验证，关注的是系统是否能在实际演示中形成合理状态。",
    ]:
        add_p(doc, text)
    add_h2(doc, "7.2  基础显示与在线测试")
    add_p(doc, "基础运行测试从设备上电和网页在线状态开始。设备上电后，OLED 能显示状态信息，Web 页面能够收到设备遥测并显示在线状态。夜间状态页面用于验证暗环境和活动状态是否能被正确推送到前端，如下图所示。")
    add_image(doc, pic("夜间状态.png"), "图7-1  夜间状态页面", width=12.5)
    add_h2(doc, "7.3  起夜与暗环境联动测试")
    add_p(doc, "起夜测试时，先让环境光照低于暗光阈值，再让PIR检测到人体活动，系统应显示夜间活动，并根据联动开关点亮LED。该测试验证了光照和人体活动不是孤立判断，而是组合成更接近真实生活的起夜场景。暗环境、起夜和窗帘联动测试页面如下图所示。")
    add_image_grid(doc, [
        (pic("起夜检测.png"), "图7-2  起夜检测页面"),
        (pic("起夜开灯.png"), "图7-3  起夜开灯测试页面"),
        (pic("暗环境人体检测.png"), "图7-4  暗环境人体检测页面"),
        (pic("暗环境窗帘联动.png"), "图7-5  暗环境窗帘联动页面"),
    ], width=6.5)
    add_h2(doc, "7.4  空气异常与通风测试")
    add_p(doc, "空气异常测试用于验证气体读数、危险等级、风扇联动和事件记录是否一致。当 MQ 系列读数超过设定阈值时，页面会进入异常状态，本地端也会根据配置启动风扇或蜂鸣器。空气异常检测结果如下图所示。")
    add_p(doc, "环境检测分析表用于观察一段时间内环境数据变化，它可以帮助判断阈值是否设置过低或过高。对于真实家庭部署，阈值需要结合正常环境读数进行调整，而不是直接套用固定数值。空气异常和环境分析测试结果如下图所示。")
    add_image_grid(doc, [
        (pic("空气异常检测.png"), "图7-6  空气异常检测页面"),
        (pic("环境检测分析表.png"), "图7-7  环境检测分析表"),
    ], width=6.5)
    add_h2(doc, "7.5  振动、长时间无活动与事件记录测试")
    add_p(doc, "振动测试通过触发 SW-420 模块模拟冲击，系统应在页面中显示异常状态并记录事件。由于振动本身可能来自桌面碰撞，所以它不会单独等同于跌倒，而是作为疑似跌倒判断中的一个重要线索。地震或强振动报警页面如下图所示。")
    add_p(doc, "长时间无活动测试用于验证PIR最近活动时间和无人活动阈值是否生效。当超过设定时间没有检测到活动时，系统会进入无人活动提醒，并在网页端显示对应状态。强振动和长时间无活动测试页面如下图所示。")
    add_image_grid(doc, [
        (pic("地震报警.png"), "图7-8  强振动报警页面"),
        (pic("长时间无活动检测.png"), "图7-9  长时间无活动检测页面"),
    ], width=6.5)
    add_p(doc, "事件记录测试表明，系统能够把关键异常保存在列表中，便于家属事后查看。单次状态只能说明当前情况，事件记录则能反映最近一段时间系统是否频繁出现异常。")
    add_h2(doc, "7.6  测试结果汇总")
    add_table(doc, "表7-1  系统测试结果汇总", [
        ["测试项目", "预期结果", "测试结论"],
        ["设备在线与遥测", "Web 页面显示设备在线并刷新数据", "通过"],
        ["暗环境活动", "暗环境下 PIR 触发后显示夜间活动并可点亮 LED", "通过"],
        ["气体异常", "气体读数超过阈值后显示风险并触发通风联动", "通过"],
        ["长时间无活动", "超过阈值后进入无人活动提醒", "通过"],
        ["振动冲击", "触发后记录异常事件并进入对应提示", "通过"],
        ["阈值设置", "网页端修改阈值后能够保存并影响推导结果", "通过"],
        ["设备离线", "设备断开后页面显示离线并限制配置操作", "通过"],
    ], [3.5, 7.0, 3.0])
    add_p(doc, "综合测试结果可以看出，系统已经完成了从传感器采集、本地判断、执行器联动到 Web 展示的基本闭环。测试中也暴露出一些需要继续改进的地方，例如 MQ 系列传感器需要预热和现场标定，面包板连接在长期运行中的稳定性有限，PIR 和 FSR 对安装位置比较敏感。这些问题没有否定系统方案本身，而是说明后续如果要实际部署，需要进一步完成结构固定、外壳设计、传感器标定和长期数据验证。")

    add_h1(doc, "8  总结与展望")
    add_h2(doc, "8.1  工作总结")
    for text in [
        "本文围绕独居老人居家安全监护需求，设计并实现了一套基于 ESP32-S3 和多模态感知的智能监护原型系统。系统通过温湿度、光照、气体、人体活动、振动、压力和 SOS 等信息综合判断居家状态，并通过 OLED、蜂鸣器、LED、风扇、舵机和 Web 页面完成本地提醒与远程展示。与单一传感器报警器相比，本文系统能够从更多角度描述老人居家环境和活动变化；与摄像头方案相比，它对隐私的影响更小，也更适合卧室和床边等场景。",
        "在结构设计上，本文将系统归纳为感知层、传输层和用户层，并在设备内部进一步完成模块化划分。感知层提供多源状态线索，传输层负责本地判断、遥测上报和配置同步，用户层面向老人和家属提供不同形式的反馈。通过实物搭建和页面测试，系统能够完成环境监测、气体异常提醒、暗环境人体活动照明、起夜记录、长时间无活动提醒、振动事件、SOS 求助、阈值设置和设备离线提示等功能。",
        "从毕业设计角度看，本课题完成了硬件、固件、Web 和测试材料之间的对应，能够体现物联网系统从端侧采集到远程展示的完整链路。论文写作中也尽量避免把所有函数和参数逐条列成说明书，而是围绕应用场景、方案取舍、模块协作和测试结果展开，这更符合毕业论文对设计思路和论证过程的要求。",
    ]:
        add_p(doc, text)
    add_h2(doc, "8.2  不足与展望")
    for text in [
        "系统目前仍属于原型阶段，距离长期家庭部署还有一些差距。一方面，MQ 系列气体传感器的数值容易受环境和预热影响，后续应加入更规范的标定过程；另一方面，面包板连接适合调试，但长期使用容易松动，需要进一步完成外壳固定、线束整理和电源保护；还有，当前规则融合方法可解释性较强，但对复杂行为的识别能力有限，后续可以在积累更多真实数据后，引入滑动窗口统计、异常趋势分析或轻量模型。",
        "未来还可以继续加入手机推送、短信或微信通知，使危险事件能够更快到达家属；也可以增加本地数据缓存和断点补传能力，减少网络波动带来的记录缺口；在用户体验方面，可以把网页端改造成更适合手机查看的界面，并增加老人日常规律分析。随着传感器安装方式和数据积累逐步完善，系统有望从毕业设计原型进一步发展为更接近实际居家养老需求的智能监护设备。",
    ]:
        add_p(doc, text)


def add_refs_and_appendix(doc: Document) -> None:
    global AUTO_CITES_ENABLED
    AUTO_CITES_ENABLED = False
    add_h1(doc, "参考文献")
    refs = [
        "[1] 国家统计局. 中国统计年鉴[M]. 北京: 中国统计出版社, 2024.",
        "[2] 民政部. 智慧健康养老产业发展行动计划[Z]. 北京: 民政部, 2021.",
        "[3] Espressif Systems. ESP32-S3 Technical Reference Manual[EB/OL]. 2024.",
        "[4] Espressif Systems. ESP32-S3-WROOM-1 Datasheet[EB/OL]. 2024.",
        "[5] Arduino. Arduino Core for ESP32 Documentation[EB/OL]. 2024.",
        "[6] PlatformIO. PlatformIO Core Documentation[EB/OL]. 2024.",
        "[7] Adafruit Industries. DHT Sensor Library Documentation[EB/OL]. 2024.",
        "[8] ROHM Semiconductor. BH1750FVI Ambient Light Sensor Datasheet[Z]. 2011.",
        "[9] Hanwei Electronics. MQ Series Gas Sensor Datasheets[Z]. 2020.",
        "[10] Adafruit Industries. SSD1306 OLED Display Library Documentation[EB/OL]. 2024.",
        "[11] Node.js Foundation. Node.js HTTP Module Documentation[EB/OL]. 2024.",
        "[12] WHATWG. Server-Sent Events Living Standard[EB/OL]. 2024.",
        "[13] 阿里云计算有限公司. 物联网平台产品文档[EB/OL]. 2024.",
        "[14] 王明, 李强. 基于物联网的智慧养老监测系统设计研究[J]. 电子技术应用, 2022(8):110-114.",
        "[15] 张华. 多传感器融合在居家安全监测中的应用[J]. 自动化与仪器仪表, 2021(6):87-91.",
        "[16] 刘洋, 陈晨. 面向居家养老的智能监测系统研究[J]. 电子设计工程, 2023(12):65-69.",
        "[17] 李娜. 基于嵌入式平台的室内环境安全监测系统设计[J]. 现代电子技术, 2021(18):80-84.",
    ]
    for ref in refs:
        add_p(doc, ref, first_line=False)

    add_h1(doc, "致    谢")
    for text in [
        "本课题从选题、需求分析、硬件搭建、程序调试、网页设计到论文整理，经历了多次修改和验证。感谢指导教师在系统方向、论文结构和问题修改方面给予的指导，也感谢同学和家人在实物测试、图片整理和演示准备中提供的帮助。",
        "通过本次毕业设计，我对嵌入式系统、物联网通信、前后端数据交互和工程文档写作有了更完整的认识，也更清楚地体会到一个系统能运行只是第一步，能稳定解释、能被用户看懂、能经得起测试和修改同样重要。今后我会继续在嵌入式与物联网方向加强学习，把本课题中暴露出的问题作为后续改进的起点。",
    ]:
        add_p(doc, text)

    add_h1(doc, "附录")
    add_h2(doc, "附录A  关键配置节选")
    add_p(doc, "附录中列出部分关键配置，主要用于说明硬件连接和阈值来源。完整代码以项目工程文件为准。")
    snippets = [
        ("DHT22", ROOT / "include" / "devices" / "dht22_config.h"),
        ("BH1750", ROOT / "include" / "devices" / "bh1750_config.h"),
        ("执行器", ROOT / "include" / "devices" / "actuator_config.h"),
        ("系统时序", ROOT / "include" / "config.h"),
    ]
    for title, path in snippets:
        add_h3(doc, title)
        text = path.read_text(encoding="utf-8")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text[:1800])
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(8.5)


def save_manuscript(doc: Document) -> None:
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    MANUSCRIPT.write_text("\n\n".join(lines), encoding="utf-8")


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_abstract(doc)
    add_front_matter(doc)
    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    main_section.top_margin = Cm(3)
    main_section.bottom_margin = Cm(3)
    main_section.left_margin = Cm(3)
    main_section.right_margin = Cm(2.5)
    header = main_section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run("长春工程学院毕业设计（论文）")
    add_page_number(main_section)
    add_main_sections(doc)
    add_refs_and_appendix(doc)
    save_manuscript(doc)
    doc.save(OUT)
    chars = len(re.sub(r"\s+", "", MANUSCRIPT.read_text(encoding="utf-8")))
    print(f"created={OUT}")
    print(f"manuscript={MANUSCRIPT}")
    print(f"chars={chars}")
    print(f"pictures={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()

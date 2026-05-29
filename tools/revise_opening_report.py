"""Revise only the technical route and main functions in the opening report."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = Path(r"C:\Users\19573\Desktop\论文\开题报待修改.docx")
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "开题报告_技术路线和主要功能修改版.docx"


REPLACEMENT_PARAGRAPHS = [
    ("label", "技术路线："),
    (
        "body",
        "本课题采用“多源感知层—ESP32-S3 主控层—本地执行与显示层—Web 可视化与配置层”的分层技术路线。"
        "多源感知层由 DHT22 温湿度、BH1750 光照、MQ2 烟雾/可燃气体、MQ135 空气质量、MQ7 一氧化碳、"
        "HC-SR501 人体红外、SW-420 振动、FSR402 压力和 SOS 按键组成，负责采集老人居家环境、活动和安全状态。"
        "ESP32-S3 主控层基于 PlatformIO 与 Arduino Framework 实现周期采样、阈值判断、状态融合、告警优先级判断和遥测数据封装。",
    ),
    (
        "body",
        "本地执行与显示层由 SSD1306 OLED、蜂鸣器、LED 灯、继电器风扇和 SG90 舵机构成。"
        "当系统检测到气体异常、温湿度异常、暗环境人体活动、长时间无活动、疑似跌倒或 SOS 求助时，"
        "可在本地完成声光提醒、通风联动、灯光联动和舵机动作，保证网络异常时仍具备基本报警能力。"
        "Web 可视化与配置层采用 Node.js 后端和网页 Dashboard，ESP32 通过 HTTP Mirror 将数据上报到 /api/telemetry，"
        "网页端通过 SSE 实时展示状态，并支持阈值设置、联动开关、事件记录、起夜记录和设备离线提示。",
    ),
    (
        "body",
        "整体流程为：传感器采集数据 → ESP32-S3 进行阈值判断和多源状态融合 → OLED 与执行器完成本地反馈 → "
        "ESP32 将标准遥测数据上报至 Web 后端 → 网页端展示实时监测结果并下发控制参数。"
        "该路线与当前系统实现一致，能够体现非接触式、多传感器融合、本地自治和远程可视化的设计特点。",
    ),
    ("label", "主要功能："),
    (
        "body",
        "1. 环境监测功能：采集温度、湿度和光照强度，判断高温、高湿、低温、低湿和暗环境状态，为通风控制和夜间照明提供依据。",
    ),
    (
        "body",
        "2. 气体安全监测功能：通过 MQ2、MQ135、MQ7 分别监测烟雾/可燃气体、空气质量和一氧化碳风险，"
        "根据预警阈值和危险阈值生成不同等级告警，并联动风扇、蜂鸣器和网页强提醒。",
    ),
    (
        "body",
        "3. 人体活动与起夜辅助功能：利用 PIR 检测人体活动，结合 BH1750 光照判断暗环境活动；"
        "当夜间检测到老人活动或离床时，自动点亮 LED 并在网页端记录起夜事件。",
    ),
    (
        "body",
        "4. 压力、振动与跌倒风险识别功能：利用 FSR402 判断床位占用和压力异常，利用 SW-420 检测振动，"
        "并结合 PIR 长时间无活动状态，对疑似跌倒和异常静止情况进行综合判断。",
    ),
    (
        "body",
        "5. SOS 主动求助功能：老人按下 SOS 按键后，系统立即触发蜂鸣器、LED、舵机动作和网页端紧急提醒，"
        "提升突发情况下的求助效率。",
    ),
    (
        "body",
        "6. Web 远程监测与配置功能：网页端展示实时数据、危险等级、设备在线状态、历史趋势和事件记录，"
        "支持传感器启用状态、阈值参数和联动开关配置，使系统便于调试、演示和后续扩展。",
    ),
]


def format_run(run, kind: str) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)
    run.bold = kind == "label"


def format_paragraph(paragraph, kind: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0 if kind == "label" else 22)
    for run in paragraph.runs:
        format_run(run, kind)


def append_paragraph_before(cell, before_element, text: str, kind: str) -> None:
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    format_run(run, kind)
    format_paragraph(paragraph, kind)
    new_element = paragraph._p
    cell._tc.remove(new_element)
    before_element.addprevious(new_element)


def revise_target_block(cell) -> None:
    paragraphs = list(cell.paragraphs)
    start = None
    end = None
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text == "技术路线：":
            start = index
        if start is not None and text == "课题可能存在的问题":
            end = index
            break

    if start is None or end is None or start >= end:
        raise RuntimeError("未找到可替换的“技术路线/主要功能”区块，已停止修改。")

    before_element = paragraphs[end]._p
    for kind, text in REPLACEMENT_PARAGRAPHS:
        append_paragraph_before(cell, before_element, text, kind)

    for paragraph in paragraphs[start:end]:
        cell._tc.remove(paragraph._p)


def main() -> None:
    document = Document(SOURCE_DOCX)
    revise_target_block(document.tables[1].cell(0, 0))
    document.core_properties.comments = "仅修改技术路线和主要功能，参考文献及其他栏目保持不变。"
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
